# ============================================================
# invoice_extractor.py
# Extension module of the AI Email Digest Agent
#
# Responsibility:
#   - Detect invoice/bill/receipt attachments in emails
#   - Extract text from PDF, Image (PNG/JPG), JSON formats
#   - Redact sensitive info before LLM processing
#   - Use Groq LLM to extract structured billing fields
#   - Append results to a dated CSV file
# ============================================================

import os
import re
import io
import json
import csv
import base64
import logging
from datetime import date
from dotenv import load_dotenv

load_dotenv(override=True)

# ──────────────────────────────────────────────
# Setup logger
# ──────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="[%(asctime)s] %(levelname)s - %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S"
)
logger = logging.getLogger("invoice_extractor")

# ──────────────────────────────────────────────
# CSV output file (dated)
# ──────────────────────────────────────────────
OUTPUT_CSV = f"invoices_extracted_{date.today().strftime('%Y-%m-%d')}.csv"

CSV_HEADERS = [
    "Vendor Name",
    "Final Total",
    "Currency",
    "Due Date",
    "Payment Status",
    "Extraction Status",
    "Invoice Number",
    "Line Items (Summary)",
    "Invoice Date",
    "Subtotal",
    "Tax",
    "Discount",
    "Vendor Email",
    "Vendor Phone",
    "Client Name",
    "Attachment Filename",
    "Source Email Link",
    "Local File Path",
    "Error Notes"
]

# ──────────────────────────────────────────────
# Trigger keywords: filenames suggesting invoices
# ──────────────────────────────────────────────
INVOICE_KEYWORDS = [
    "invoice", "bill", "receipt", "payment", "statement",
    "order", "purchase", "transaction", "proforma", "tax_invoice",
    "utility", "salary", "slip", "gst", "vat", "service", "booking", "summary"
]

SUPPORTED_EXTENSIONS = {
    "pdf":  "pdf",
    "png":  "image",
    "jpg":  "image",
    "jpeg": "image",
    "csv":  "csv",
    "doc":  "word",
    "docx": "word"
}


# ============================================================
# STEP 1: Check if an attachment is likely an invoice
# ============================================================
def is_invoice_attachment(filename: str, email_subject: str = "") -> bool:
    name_lower = filename.lower()
    ext = name_lower.rsplit(".", 1)[-1] if "." in name_lower else ""
    if ext not in SUPPORTED_EXTENSIONS:
        return False
    
    # Check if filename contains trigger keywords
    if any(kw in name_lower for kw in INVOICE_KEYWORDS):
        return True
        
    # Fallback: if filename is generic (like "document.pdf") but subject says "invoice"
    subject_lower = email_subject.lower()
    if any(kw in subject_lower for kw in INVOICE_KEYWORDS):
        return True
        
    return False


# ============================================================
# STEP 2: Download attachment bytes from Gmail
# ============================================================
def download_attachment(service, message_id: str, attachment_id: str) -> bytes:
    attachment = service.users().messages().attachments().get(
        userId="me",
        messageId=message_id,
        id=attachment_id
    ).execute()
    data = attachment.get("data", "")
    return base64.urlsafe_b64decode(data)


# ============================================================
# STEP 3A: Extract text from PDF (pdfplumber → OCR fallback)
# ============================================================
def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        text = "\n".join(text_parts).strip()
        if text:
            return text
        # If no text extracted (scanned PDF), fall back to OCR
        logger.info("PDF has no embedded text — falling back to OCR.")
        return _ocr_pdf_pages(file_bytes)
    except Exception as e:
        logger.warning(f"PDF extraction failed: {e}")
        return ""


def _ocr_pdf_pages(file_bytes: bytes) -> str:
    """Convert each PDF page to image and OCR it natively without Tesseract."""
    try:
        import pdfplumber
        import numpy as np
        from PIL import Image
        from rapidocr_onnxruntime import RapidOCR
        
        ocr = RapidOCR()
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                # Get numpy array directly instead of PIL Image for rapidocr
                img = page.to_image(resolution=200).original
                img_array = np.array(img)
                
                result, _ = ocr(img_array)
                if result:
                    # result is a list of tuples: (box, text, score)
                    page_text = "\n".join([line[1] for line in result])
                    text_parts.append(page_text)
                    
        return "\n".join(text_parts).strip()
    except Exception as e:
        logger.warning(f"OCR on PDF failed: {e}")
        return ""


# ============================================================
# STEP 3B: Extract text from image using pytesseract
# ============================================================
def extract_text_from_image(file_bytes: bytes) -> str:
    try:
        import numpy as np
        from PIL import Image
        from rapidocr_onnxruntime import RapidOCR

        ocr = RapidOCR()
        
        # Open bytes as PIL Image, convert to RGB, then to numpy array for RapidOCR
        img = Image.open(io.BytesIO(file_bytes)).convert('RGB')
        img_array = np.array(img)
        
        result, _ = ocr(img_array)
        if result:
            return "\n".join([line[1] for line in result]).strip()
        return ""
    except Exception as e:
        logger.warning(f"Image OCR failed: {e}")
        return ""


# ============================================================
# STEP 3C: Extract text from CSV attachment
# ============================================================
def extract_text_from_csv(file_bytes: bytes) -> str:
    try:
        text = file_bytes.decode("utf-8", errors="replace")
        return text.strip()
    except Exception as e:
        logger.warning(f"CSV extraction failed: {e}")
        return ""

# ============================================================
# STEP 3D: Extract text from Word attachment
# ============================================================
def extract_text_from_word(file_bytes: bytes) -> str:
    try:
        import docx
        import io
        doc = docx.Document(io.BytesIO(file_bytes))
        fullText = []
        
        # 1. Extract from standard paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                fullText.append(para.text.strip())
        
        # 2. Extract from tables (Crucial for invoices!)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Collect text from paragraphs inside the cell
                    for para in cell.paragraphs:
                        if para.text.strip():
                            fullText.append(para.text.strip())
        
        return "\n".join(fullText).strip()
    except ImportError:
        logger.error("python-docx is not installed. Please install it using 'pip install python-docx'.")
        return ""
    except Exception as e:
        logger.warning(f"Word extraction failed: {e}")
        return ""


# ============================================================
# STEP 4: Route to correct extractor based on file extension
# ============================================================
def extract_text_from_attachment(filename: str, file_bytes: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    fmt = SUPPORTED_EXTENSIONS.get(ext, "unknown")

    if fmt == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif fmt == "image":
        return extract_text_from_image(file_bytes)
    elif fmt == "csv":
        return extract_text_from_csv(file_bytes)
    elif fmt == "word":
        return extract_text_from_word(file_bytes)
    else:
        logger.warning(f"Unsupported format for: {filename}")
        return ""


# ============================================================
# STEP 5: Redact sensitive data before sending to LLM
# (inherits same rules as parent agent)
# ============================================================
def redact_sensitive_info(text: str) -> str:
    # Phone numbers (escaping hyphens in character classes to avoid range errors)
    text = re.sub(
        r'(\+?\d{1,3}[.\-\s]?)?(\(?\d{3}\)?[.\-\s]?)\d{3}[.\-\s]?\d{4}',
        'REDACTED_PHONE', text
    )
    # Email addresses
    text = re.sub(
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b',
        'REDACTED_EMAIL', text
    )
    # Credit card numbers (13-16 digits)
    text = re.sub(r'\b(?:\d{4}[ -]?){3}\d{4}\b', 'REDACTED_CARD', text)
    # UPI IDs
    text = re.sub(r'\b[a-zA-Z0-9._-]+@[a-zA-Z]{3,}\b', 'REDACTED_UPI', text)
    # Bank account numbers (8–12 digits)
    text = re.sub(r'\b\d{8,12}\b', 'REDACTED_ACCOUNT', text)
    # IFSC codes
    text = re.sub(r'\b[A-Z]{4}0[A-Z0-9]{6}\b', 'REDACTED_IFSC', text)
    return text


# ============================================================
# STEP 6: Send redacted text to LLM, extract structured fields
# ============================================================
def extract_invoice_fields_via_llm(text: str, filename: str) -> dict:
    from langchain_groq import ChatGroq

    groq_api_key = os.getenv("GROQ_API_KEY")
    if not groq_api_key:
        raise ValueError("GROQ_API_KEY not found!")

    llm = ChatGroq(model="llama-3.3-70b-versatile", api_key=groq_api_key)

    prompt = f"""You are a financial data extraction specialist.
Extract the following fields from the invoice text below.
If a field is missing, use exactly "N/A".

Return ONLY a valid raw JSON object (no markdown, no explanation) with these exact keys:
{{
  "invoice_number": "",
  "invoice_date": "",
  "vendor_name": "",
  "vendor_email": "",
  "vendor_phone": "",
  "client_name": "",
  "line_items": "",
  "subtotal": "",
  "tax": "",
  "discount": "",
  "final_total": "",
  "currency": "",
  "due_date": "",
  "payment_status": ""
}}

Rules:
- "line_items" should be a neat, numbered list using newline characters (\\n). e.g.:
   1. Item Name - Qty: X - Unit Price: $Y - Total: $Z
   2. Item Name - Qty: X - Unit Price: $Y - Total: $Z
- All monetary fields (subtotal, tax, discount, final_total) MUST include the currency symbol (e.g., $150.00, €30.00) and proper decimal formatting.
- Sensitive fields (card numbers, account numbers) already appear as REDACTED_* tokens — keep them as-is
- Do NOT add fields not listed above

Invoice Text (from file: {filename}):
\"\"\"
{text[:4000]}
\"\"\"
"""

    response = llm.invoke(prompt)
    raw = response.content if hasattr(response, "content") else str(response)

    # Clean up possible markdown code fences
    raw = re.sub(r"```(?:json)?", "", raw).strip().strip("`").strip()

    try:
        fields = json.loads(raw)
    except json.JSONDecodeError:
        logger.warning(f"LLM returned non-JSON for {filename}. Attempting partial parse.")
        fields = _fallback_parse(raw)

    return fields


def _fallback_parse(raw: str) -> dict:
    """Try to extract JSON from LLM response even if it's partially wrapped."""
    match = re.search(r'\{.*\}', raw, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except Exception:
            pass
    return {}


# ============================================================
# STEP 7: Write one row to the dated CSV
# ============================================================
def append_to_csv(row: dict):
    global OUTPUT_CSV
    
    def try_write(filename):
        file_exists = os.path.exists(filename)
        with open(filename, mode="a", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=CSV_HEADERS)
            if not file_exists:
                writer.writeheader()
            writer.writerow(row)
        return filename

    try:
        saved_path = try_write(OUTPUT_CSV)
        logger.info(f"✅ Invoice row saved to {saved_path}")
    except PermissionError:
        # If locked, try a secondary file name so the data isn't lost
        alt_name = OUTPUT_CSV.replace(".csv", "_v2.csv")
        logger.warning(f"⚠️ {OUTPUT_CSV} is LOCKED (open in Excel?). Trying {alt_name} instead...")
        try:
            saved_path = try_write(alt_name)
            logger.info(f"✅ Invoice row saved to fallback file: {saved_path}")
        except PermissionError:
            logger.error("❌ CRITICAL: Both primary and fallback CSV files are LOCKED. "
                         "Close your Excel/CSV viewers and try again.")
            raise
    except Exception as e:
        logger.error(f"❌ Failed to write to CSV: {e}")
        raise


# ============================================================
# MAIN ENTRY POINT — called from main pipeline
# ============================================================
def process_single_attachment(service, message_id: str, email_subject: str, filename: str, attachment_id: str = None, file_bytes: bytes = None):
    """
    Extracted logic to process a specific attachment.
    Can be called with attachment_id (for Gmail) OR file_bytes (for manual upload).
    """
    # Resolve the local save path upfront so it's available in both success and failure branches
    static_invoice_dir = os.path.join("static", "processed_invoices")
    os.makedirs(static_invoice_dir, exist_ok=True)
    safe_filename = f"{message_id}_{filename}"
    local_path = os.path.join(static_invoice_dir, safe_filename)
    local_url = f"/processed_invoices/{safe_filename}"

    try:
        if not file_bytes and attachment_id:
            # Download raw bytes from Gmail
            file_bytes = download_attachment(service, message_id, attachment_id)

        if not file_bytes:
            raise ValueError("No file content provided.")

        # Save a local copy immediately so the file is always available for preview/download,
        # even if LLM extraction later fails
        with open(local_path, "wb") as f:
            f.write(file_bytes)

        # Extract text
        raw_text = extract_text_from_attachment(filename, file_bytes)
        if not raw_text.strip():
            raise ValueError("No text extracted from attachment.")

        # Redact before sending to LLM
        redacted_text = redact_sensitive_info(raw_text)

        # LLM extraction
        fields = extract_invoice_fields_via_llm(redacted_text, filename)

        # Generate a clickable Gmail link (only if it's from an email)
        gmail_link = f"https://mail.google.com/mail/u/0/#inbox/{message_id}" if message_id != "MANUAL_UPLOAD" else "N/A"

        row = {
            "Vendor Name":         fields.get("vendor_name",     "N/A"),
            "Final Total":         fields.get("final_total",     "N/A"),
            "Currency":            fields.get("currency",        "N/A"),
            "Due Date":            fields.get("due_date",        "N/A"),
            "Payment Status":      fields.get("payment_status",  "N/A"),
            "Extraction Status":   "SUCCESS",
            "Invoice Number":      fields.get("invoice_number",  "N/A"),
            "Line Items (Summary)":fields.get("line_items",      "N/A"),
            "Invoice Date":        fields.get("invoice_date",    "N/A"),
            "Subtotal":            fields.get("subtotal",        "N/A"),
            "Tax":                 fields.get("tax",             "N/A"),
            "Discount":            fields.get("discount",        "N/A"),
            "Vendor Email":        fields.get("vendor_email",    "N/A"),
            "Vendor Phone":        fields.get("vendor_phone",    "N/A"),
            "Client Name":         fields.get("client_name",     "N/A"),
            "Attachment Filename": filename,
            "Source Email Link":   gmail_link,
            "Local File Path":     local_url,
            "Error Notes":         ""
        }

    except Exception as e:
        logger.error(f"❌ Failed to process {filename}: {e}")
        gmail_link = f"https://mail.google.com/mail/u/0/#inbox/{message_id}" if message_id != "MANUAL_UPLOAD" else "N/A"
        row = {
            "Vendor Name":         "N/A",
            "Final Total":         "N/A",
            "Currency":            "N/A",
            "Due Date":            "N/A",
            "Payment Status":      "N/A",
            "Extraction Status":   "FAILED",
            "Invoice Number":      "N/A",
            "Line Items (Summary)":"N/A",
            "Invoice Date":        "N/A",
            "Subtotal":            "N/A",
            "Tax":                 "N/A",
            "Discount":            "N/A",
            "Vendor Email":        "N/A",
            "Vendor Phone":        "N/A",
            "Client Name":         "N/A",
            "Attachment Filename": filename,
            "Source Email Link":   gmail_link,
            "Local File Path":     local_url if os.path.exists(local_path) else "",
            "Error Notes":         str(e)
        }

    append_to_csv(row)
    return row

def process_email_for_invoices(service, message_id: str, email_subject: str, parts: list):
    """
    Scan a single email's MIME parts for invoice attachments.
    """
    found_any = False

    def recurse_parts(parts_list):
        nonlocal found_any
        for part in parts_list:
            if part.get("mimeType", "").startswith("multipart"):
                recurse_parts(part.get("parts", []))
                continue

            filename = part.get("filename", "")
            if not filename or not is_invoice_attachment(filename, email_subject):
                continue

            found_any = True
            attachment_id = part.get("body", {}).get("attachmentId")
            if not attachment_id: continue

            logger.info(f"📎 Invoice attachment detected: {filename} (email: {email_subject!r})")
            process_single_attachment(service, message_id, email_subject, filename, attachment_id=attachment_id)

    recurse_parts(parts)
    if not found_any:
        logger.info(f"No invoice attachments found in message {message_id}.")
